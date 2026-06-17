from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(r"C:\Users\25124968R\Documents\Codex\2026-06-14\octa-c-users-25124968r-desktop-cuhk")
OUT = BASE / "outputs"
DOCX_OUT = OUT / "supplementary_octa_cho_table_design.docx"
MD_OUT = OUT / "supplementary_octa_cho_table_design.md"


def read_outputs() -> dict[str, pd.DataFrame]:
    return {
        "analysis": pd.read_csv(OUT / "R_repro_analysis_dataset.csv"),
        "octa_qc": pd.read_csv(OUT / "R_repro_octa_qc_summary.csv"),
        "cho_qc": pd.read_csv(OUT / "R_repro_cho_qc_summary.csv"),
        "perf": pd.read_csv(OUT / "R_repro_auc_performance.csv"),
        "diff": pd.read_csv(OUT / "R_repro_auc_pairwise_differences.csv"),
        "overview": pd.read_csv(OUT / "meeting_auc_overview.csv"),
        "keydiff": pd.read_csv(OUT / "meeting_auc_key_differences.csv"),
    }


def fmt_median_iqr(x: pd.Series, digits: int = 2) -> str:
    x = pd.to_numeric(x, errors="coerce").dropna()
    if len(x) == 0:
        return ""
    q1 = x.quantile(0.25)
    med = x.quantile(0.50)
    q3 = x.quantile(0.75)
    return f"{med:.{digits}f} ({q1:.{digits}f}, {q3:.{digits}f})"


def fmt_n_pct(mask: pd.Series, denom: int) -> str:
    n = int(mask.sum())
    pct = 100 * n / denom if denom else 0
    return f"{n} ({pct:.1f})"


def cohort_table(df: pd.DataFrame) -> list[list[str]]:
    groups = {
        "Total (n = 137)": df,
        "Non-myopia (n = 22)": df[df["SE"] > -0.50],
        "Moderate myopia (n = 91)": df[(df["SE"] > -6.00) & (df["SE"] <= -0.50)],
        "High myopia (n = 24)": df[df["SE"] <= -6.00],
    }

    rows = [["Variables", *groups.keys()]]
    rows.append(["Age, median (Q1, Q3)", *[fmt_median_iqr(g["Age_canonical"], 1) for g in groups.values()]])
    rows.append(["Sex = 1, n (%)", *[fmt_n_pct(g["Sex_canonical"] == 1, len(g)) for g in groups.values()]])
    rows.append(["Sex = 2, n (%)", *[fmt_n_pct(g["Sex_canonical"] == 2, len(g)) for g in groups.values()]])
    rows.append(["Axial length, median (Q1, Q3), mm", *[fmt_median_iqr(g["AL"], 2) for g in groups.values()]])
    rows.append(["Spherical equivalent, median (Q1, Q3), D", *[fmt_median_iqr(g["SE"], 2) for g in groups.values()]])
    rows.append(["AL >= 26 mm, n (%)", *[fmt_n_pct(g["AL"] >= 26.0, len(g)) for g in groups.values()]])
    return rows


def build_tables(data: dict[str, pd.DataFrame]) -> list[dict]:
    analysis = data["analysis"]
    octa_qc = data["octa_qc"]
    cho_qc = data["cho_qc"]
    overview = data["overview"]
    keydiff = data["keydiff"]

    table1 = [
        ["Domain", "Module / biomarker group", "Components used for score", "Interpretation", "Used in model(s)"],
        ["OCTA", "CCFA", "Median of retained choriocapillaris flow-area grid values", "Flow-related choriocapillaris signal", "Model 3, Model 5"],
        ["OCTA", "CT", "Median of retained choroidal thickness grid values", "Structural choroidal thickness", "Model 3, Model 4, Model 5"],
        ["OCTA", "CVI", "Median of retained choroidal vascularity index grid values", "Relative choroidal vascularity", "Model 3, Model 4, Model 5"],
        ["Cho/CFP", "Density", "Whole-region GAN-derived vessel-density metrics", "Global choroidal vessel density", "Model 2, Model 5"],
        ["Cho/CFP", "Complexity", "Whole-region GAN-derived complexity metrics", "Branching/structural complexity", "Model 2, Model 5"],
        ["Cho/CFP", "Calibre", "Whole-region GAN-derived calibre metrics", "Vessel width/calibre profile", "Model 2, Model 5"],
        ["Cho/CFP", "Tortuosity", "Whole-region GAN-derived tortuosity metrics", "Curvature and vessel path irregularity", "Model 2, Model 5"],
        ["Cho/CFP", "Branching Angle", "Whole-region GAN-derived branching-angle metrics", "Angular branching geometry", "Model 2, Model 5"],
    ]

    table2 = cohort_table(analysis)

    table3 = [["Source", "Module / category", "Candidate n", "Retained n", "Removed variables", "Processing note"]]
    for _, r in octa_qc.iterrows():
        module = str(r["module"])
        retained = int(r["retained_columns"])
        removed = f"missing >90%: {int(r['removed_missing_rate_gt_90pct'])}; low IQR: {int(r['removed_low_iqr'])}"
        note = "Grid-level QC -> module median -> median imputation if needed -> scale"
        if module == "CFA":
            note = "QC checked; excluded from final models because of high correlation with CCFA"
        table3.append(["OCTA", module, str(int(r["input_grid_columns"])), str(retained), removed, note])
    for _, r in cho_qc.iterrows():
        removed = (
            f"missing >90%: {int(r['removed_missing_rate_gt_90pct'])}; "
            f"valid n<threshold: {int(r['removed_valid_n_lt_threshold'])}; "
            f"low IQR: {int(r['removed_low_iqr'])}"
        )
        table3.append([
            "Cho/CFP",
            str(r["category"]),
            str(int(r["matched_variables"])),
            str(int(r["retained_variables"])),
            removed,
            "Variable median imputation -> variable z-score -> module median -> scale",
        ])

    table4 = [
        ["Section", "Name", "Definition / predictors", "Event count"],
        ["Outcome", "High myopia", "SE <= -6.00 D", "24/137"],
        ["Outcome", "Moderate myopia", "-6.00 D < SE <= -0.50 D", "91/137"],
        ["Outcome", "Long axial length", "AL >= 26 mm", "25/137"],
        ["Model", "Model 1 Clinical", "Age + Sex", ""],
        ["Model", "Model 2 Clinical + Cho", "Age + Sex + 5 Cho/CFP modules", ""],
        ["Model", "Model 3 Clinical + OCTA", "Age + Sex + CCFA + CT + CVI", ""],
        ["Model", "Model 4 Clinical + OCTA_nonflow", "Age + Sex + CT + CVI", ""],
        ["Model", "Model 5 Clinical + OCTA + Cho", "Age + Sex + OCTA modules + 5 Cho/CFP modules", ""],
    ]

    # Keep the same compact presentation table already approved for meeting use.
    table5 = [["Endpoint", "Events / N", "Clinical", "+ Cho", "+ OCTA", "+ OCTA_nonflow", "+ OCTA + Cho"]]
    for _, r in overview.iterrows():
        table5.append([
            str(r["Endpoint"]),
            str(r["Events / N"]),
            f"{float(r['M1 Clinical']):.3f}",
            f"{float(r['M2 + Cho']):.3f}",
            f"{float(r['M3 + OCTA']):.3f}",
            f"{float(r['M4 + OCTA_nonflow']):.3f}",
            f"{float(r['M5 + OCTA + Cho']):.3f}",
        ])

    keep_comparisons = {
        "Cho vs OCTA",
        "Cho vs OCTA_nonflow",
        "OCTA flow contribution",
        "Add Cho beyond OCTA",
    }
    table6 = [["Endpoint", "Comparison", "Delta CV AUC", "95% CI", "p", "Interpretation"]]
    for _, r in keydiff[keydiff["Comparison"].isin(keep_comparisons)].iterrows():
        table6.append([
            str(r["Endpoint"]),
            str(r["Comparison"]),
            f"{float(r['Delta CV AUC']):.3f}",
            str(r["95% CI"]),
            str(r["p"]),
            str(r["Interpretation"]),
        ])

    return [
        {
            "caption": "Supplementary Table 1. Definition of OCTA and Cho/CFP biomarker modules.",
            "rows": table1,
            "note": "CCFA, choriocapillaris flow area; CT, choroidal thickness; CVI, choroidal vascularity index; CFP, color fundus photography.",
            "widths": [1000, 1700, 2600, 2300, 1760],
        },
        {
            "caption": "Supplementary Table 2. Characteristics of the final OD analytic cohort stratified by refractive status.",
            "rows": table2,
            "note": "Sex was kept as the coded variable available in the source files. Moderate and high myopia are mutually exclusive by SE definition.",
            "widths": [2600, 1700, 1700, 1700, 1700],
        },
        {
            "caption": "Supplementary Table 3. Quality control and module construction summary.",
            "rows": table3,
            "note": "Both data sources removed variables with missing rate >90% and near-zero IQR. Cho/CFP additionally required valid n >= 20 or >=10% of total n.",
            "widths": [900, 1500, 1050, 950, 2250, 2710],
        },
        {
            "caption": "Supplementary Table 4. Binary outcomes and prediction model specifications.",
            "rows": table4,
            "note": "All models used the final aligned OD analysis set (n=137). Age, Sex, AL and SE were taken from the OCTA file as the canonical clinical/outcome source.",
            "widths": [1000, 2200, 4800, 1360],
        },
        {
            "caption": "Supplementary Table 5. Cross-validated AUC performance of the five prediction models.",
            "rows": table5,
            "note": "Values are repeated 5-fold cross-validated AUCs. In-sample AUCs were not prioritized for interpretation because they are more optimistic.",
            "widths": [2500, 1100, 950, 950, 950, 1350, 1560],
        },
        {
            "caption": "Supplementary Table 6. Key pairwise differences in cross-validated AUC.",
            "rows": table6,
            "note": "Delta CV AUC equals the first model minus the second model. Confidence intervals and p values were estimated using paired bootstrap on repeated-CV probabilities.",
            "widths": [2100, 1950, 1050, 1450, 750, 2060],
        },
    ]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 8.7, center: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Note. " + text)
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(80, 80, 80)


def add_supp_table(doc: Document, spec: dict) -> None:
    add_caption(doc, spec["caption"])
    rows = spec["rows"]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, spec["widths"])

    center_cols = set()
    if "AUC" in spec["caption"] or "Characteristics" in spec["caption"]:
        center_cols = set(range(1, len(rows[0])))
    if "Quality control" in spec["caption"]:
        center_cols = {2, 3}
    if "Binary outcomes" in spec["caption"]:
        center_cols = {0, 3}

    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                set_cell_text(cell, text, bold=True, font_size=8.6, center=(c_idx in center_cols))
            else:
                set_cell_text(cell, text, bold=False, font_size=8.2 if len(rows[0]) >= 6 else 8.7, center=(c_idx in center_cols))
    add_note(doc, spec["note"])


def make_docx() -> None:
    data = read_outputs()
    table_specs = build_tables(data)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run("Supplementary Tables for OCTA and Cho/CFP Biomarker Modelling")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(31, 77, 120)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    intro_run = intro.add_run(
        "Draft table set designed with the same supplement logic as the reference file: module definitions, cohort characteristics, measurement QC, model specifications, and AUC results."
    )
    intro_run.font.size = Pt(10)

    for idx, spec in enumerate(table_specs, start=1):
        if idx in {3, 5}:
            doc.add_page_break()
        add_supp_table(doc, spec)

    doc.save(DOCX_OUT)

    md_lines = ["# Supplementary Table Design", ""]
    for spec in table_specs:
        md_lines.append(f"## {spec['caption']}")
        md_lines.append("")
        rows = spec["rows"]
        widths = [max(len(str(row[i])) for row in rows) for i in range(len(rows[0]))]
        md_lines.append("| " + " | ".join(str(c).ljust(widths[i]) for i, c in enumerate(rows[0])) + " |")
        md_lines.append("| " + " | ".join("-" * widths[i] for i in range(len(widths))) + " |")
        for row in rows[1:]:
            md_lines.append("| " + " | ".join(str(c).ljust(widths[i]) for i, c in enumerate(row)) + " |")
        md_lines.append("")
        md_lines.append("Note. " + spec["note"])
        md_lines.append("")
    MD_OUT.write_text("\n".join(md_lines), encoding="utf-8")


if __name__ == "__main__":
    make_docx()
    print(DOCX_OUT)
    print(MD_OUT)
